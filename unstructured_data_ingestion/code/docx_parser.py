from docx import Document

file_path = "unstructured_data_ingestion\\data\\ZS SOW - Project Diamond - OneCI enhancements.docx"

doc = Document(file_path)

distinct_styles = set()
distinct_size = set()

for para in doc.paragraphs:
    # print(para.text)
    # print(repr(para.text))
    text = para.text.strip()

    if not text:
        continue

    style = para.style.name if para.style else "No Style"
    
    # distinct_styles.add(style)

    has_bold = False
    has_italic = False
    has_underline = False
    font_size = []

    for run in para.runs:
        if run.bold:
            has_bold = True
        if run.italic:
            has_italic = True
        if run.underline:
            has_underline = True
        if run.font.size:
            font_size.append(run.font.size.pt)

    
    distinct_size.add(tuple(font_size)) 
    print("text - ", text)
    print("style - ", style)
    print("has_bold - ", has_bold)
    print("has_italic - ", has_italic)
    print("has_underline - ", has_underline)
    print("font_size - ", font_size if font_size else "No Font Size")
    print("-" * 50)

print("Distinct Font Sizes Used in the Document:", distinct_size)

# print("Distinct Styles Used in the Document:", distinct_styles)
    # if para.text.strip():  # Check if the paragraph is not empty
    #     print(para.text)


# print(doc.paragraphs[0].text)

# print("Document Loaded Successfully!")
# print(type(doc))

